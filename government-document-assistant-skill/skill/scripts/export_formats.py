#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentExporter:
    def __init__(self):
        self.supported_formats = ['md', 'docx', 'pdf']
    
    def export_markdown(self, content: str, output_path: str) -> bool:
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f'Markdown导出失败：{e}')
            return False
    
    def export_word(self, content: str, title: str, output_path: str) -> bool:
        if not DOCX_AVAILABLE:
            print('未安装python-docx库，无法导出Word格式')
            return False
        
        try:
            doc = Document()
            
            title_para = doc.add_paragraph(title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title_run.font.name = '方正小标宋简体'
            
            for line in content.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line)
                    para_run = para.runs[0]
                    para_run.font.size = Pt(16)
                    para_run.font.name = '仿宋_GB2312'
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f'Word导出失败：{e}')
            return False
    
    def export(self, content: str, format_type: str, output_path: str, title: str = '') -> Dict[str, Any]:
        result = {
            'success': False,
            'format': format_type,
            'output_path': output_path
        }
        
        if format_type not in self.supported_formats:
            result['error'] = f'不支持的格式：{format_type}'
            return result
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        if format_type == 'md':
            result['success'] = self.export_markdown(content, output_path)
        elif format_type == 'docx':
            result['success'] = self.export_word(content, title, output_path)
        elif format_type == 'pdf':
            result['error'] = 'PDF导出功能开发中，建议先导出Word再转换为PDF'
        
        return result

def main():
    if len(sys.argv) < 4:
        print(json.dumps({'error': '参数不足：需要内容、格式、输出路径'}, ensure_ascii=False))
        sys.exit(1)
    
    content = sys.argv[1]
    format_type = sys.argv[2]
    output_path = sys.argv[3]
    title = sys.argv[4] if len(sys.argv) > 4 else ''
    
    exporter = DocumentExporter()
    result = exporter.export(content, format_type, output_path, title)
    
    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == '__main__':
    main()